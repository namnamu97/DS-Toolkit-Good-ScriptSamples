# common library
import tkinter as tk
import pickle
from pathlib import Path

# from main_command import run_word, check_read_file, delete_file

#######################################
# Command function
#######################################
# open-source lib
import docx
import pandas as pd


# standard lib
import os
import glob
import logging
from datetime import datetime
from dateutil.relativedelta import relativedelta

# config
logging.basicConfig(
    format='%(asctime)s :: %(name)s :: %(funcName)s :: Line %(lineno)d :: %(message)s',
    datefmt='%d-%m-%y %H:%M:%S',
    level = logging.INFO
    )
logger = logging.getLogger(__name__)

CWD = Path.cwd()
WORD_DIR = CWD / 'word'
Path(WORD_DIR).mkdir(parents = True, exist_ok= True)
WORD_DIR = str(CWD / WORD_DIR)
EXCEL_FILE = CWD / 'data.xlsx'
PERSIS_FILE = CWD  / 'PERSIS_FILE.pkl'

def _gen_persis_var():
    if PERSIS_FILE.is_file():
        with open(PERSIS_FILE, 'rb') as fp:
            return pickle.load(fp)
    return []

def get_code_table(file):
    logger.info(f'Get Project Code and Tables from {file}')
    doc = docx.Document(file)

    result_para = [p.text for p in doc.paragraphs]

    ma_vv = result_para[1].split(':')[1].strip()

    result_table = []

    table_dct = {}

    for c,t in enumerate(doc.tables):
        ls_table = []
        for row in t.rows:
            ls_row = []
            for cell in row.cells:
                # print(cell.text)
                ls_row.append(cell.text)
            ls_table.append(ls_row)

        if c==0:
            table_dct['General Information'] = ls_table
        elif c==1:
            table_dct['Resources Information'] = ls_table
        elif c==2:
            table_dct['CMC Interfaces'] = ls_table
        else:
            table_dct['Customer Interfaces'] = ls_table
        
    return ma_vv, table_dct

def extract_data(ma_vv, table_dct):
    logger.info('Extracting data')
    extract_dct = {}

    extract_dct['Ma VV'] = ma_vv

    for ele in table_dct['General Information']:
        if ele[0].lower() == 'project scale':
            extract_dct['Project Scale'] = ele[1]
        if ele[0].lower() == 'project cost':
            cost = int(
                ele[1]
                .replace('\n', '')
                .replace('VNĐ','')
                .replace(',','').strip()
                )
            extract_dct['Project Cost'] = cost
        if ele[0].lower() == 'project code':
            extract_dct['Project Code'] = ele[1]
        if ele[0].lower() == 'start date':
            start_date = ele[1]
            end_date = ele[3]
            extract_dct['Start Date'] = start_date
            extract_dct['End Date'] = end_date

    extract_dct['Resource'] = []
    for i,ele in enumerate(table_dct['Resources Information']):
        resource_dct = {}
        if i == 0:
            continue
        resource_dct['No'] = ele[0]
        resource_dct['Name'] = ele[1]
        resource_dct['Position'] = ele[2]
        resource_dct['Starting date'] = ele[3]
        resource_dct['Ending date'] = ele[4]

        extract_dct['Resource'].append(resource_dct)
        
    return extract_dct

def _adding_month(df):
    for i, row in df.iterrows():
        start_date = datetime.strptime(df.loc[i,'Resource_Starting date'], '%d/%m/%Y')
        end_date = datetime.strptime(df.loc[i,'Resource_Ending date'], '%d/%m/%Y')
        month_diff = (end_date.year - start_date.year) * 12 + end_date.month - start_date.month
        
        df.loc[i, f'Month 1'] = str(start_date.month)
        
        for num_month in range(2,month_diff+1+1):
            _date = start_date + relativedelta(months = num_month-1)
            df.loc[i,f'Month {num_month}'] = str(_date.month)
        
    return df


def get_dataframe(extract_dct):
    logger.info('Getting dataframe from the extracted data')
    df = pd.json_normalize(
        data = extract_dct,
        meta = [col for col in list(extract_dct.keys())[:-1]],
        record_path = 'Resource', record_prefix = 'Resource_')
    
    # change column order 
    cols = df.columns.tolist()
    cols = cols[-6:] + cols[:5]
    
    df = df[cols]

    # adding 18 Man Month col
    for i in range(1,19):
        df[f'Month {i}'] = ''

    df = _adding_month(df)

    month_dct = {
        '1':'Jan',
        '2':'Feb',
        '3':'Mar',
        '4':'Apr',
        '5':'May',
        '6':'Jun',
        '7':'Jul',
        '8':'Aug',
        '9':'Sep',
        '10':'Oct',
        '11':'Nov',
        '12':'Dec',
        '':''
    }
    
    for i in range(1,19):
        df[f'Month {i}'] = df[f'Month {i}'].map(month_dct)
        
    return df

def export_to_excel(df):
    if not EXCEL_FILE.is_file():
        df.to_excel(EXCEL_FILE, index = False)
    else:
        with pd.ExcelWriter(EXCEL_FILE, if_sheet_exists='overlay', mode = 'a', engine="openpyxl") as writer:
            startrow = writer.sheets['Sheet1'].max_row
            df.to_excel(writer, index = False, startrow = startrow, header = False)
    logger.info(f'Exported to {EXCEL_FILE}')
    
def run_word():
    PERSIS_LIST = _gen_persis_var()
    new_extracted_file = []
    for file in glob.glob(WORD_DIR + '/*.docx'):
        if file in PERSIS_LIST:
            logger.info(f'{file} is already processed')
            continue
        print(file)

        ma_vv, extract_dct = get_code_table(file)
        extract_dct = extract_data(ma_vv, extract_dct)
        df = get_dataframe(extract_dct)
        print(df)
        export_to_excel(df)

        PERSIS_LIST.append(file)
        new_extracted_file.append(file)

        with open(PERSIS_FILE, 'wb') as fp:
            pickle.dump(PERSIS_LIST, fp)

        print('-'*50)
    
    if new_extracted_file:
        text = 'Trích xuất dữ liệu thành công! Các file đã được chạy: \n\t'
        text = text + ' \n\t'.join(new_extracted_file)
    else:
        text = 'Không có dữ liệu mới được trích xuất'
    return text

def check_read_file():
    PERSIS_LIST = _gen_persis_var()
    if PERSIS_LIST:
        text = 'Các file đã được trích xuất và xử lý: \n\t'
        text = text + ' \n\t'.join(PERSIS_LIST)
        return text
    else:
        return 'Chưa có file word nào được đọc'
    # if _PERSIS_LIST:

    #     return text
    

def delete_file():
    if EXCEL_FILE.is_file() or PERSIS_FILE.is_file():
        if EXCEL_FILE.is_file():
            EXCEL_FILE.unlink()
        if PERSIS_FILE.is_file():
            PERSIS_FILE.unlink()
        return 'Các file dữ liệu đã được xóa \nBạn có thể bắt đầu trích xuất lại dữ liệu từ đầu'
    return 'Đang không có file lưu trữ dữ liệu đã được trích xuất'


#######################################
# GUI setting 
#######################################

TITLE = 'Chạy dữ liệu Word Order'

inft= tk.Tk()
inft.title(TITLE)
inft.geometry('700x600')

canvas1 = tk.Canvas(inft, width = 400, height = 250)
canvas1.pack()

# label
label = tk.Label(inft, text = TITLE, fg = 'black', font = ('Arial', 50))

# TODO: list of feature to implement

string_var = tk.StringVar()

# button location config
button_x = 150

dy_button = 35

button1_y = 20
button2_y = button1_y + dy_button
button3_y = button2_y + dy_button

# Create "Chạy dữ liệu"


def button_run_word():
    text = run_word()
    string_var.set(text)
    label = tk.Label(inft, textvariable = string_var, fg='black', font=('helvetica', 12, 'bold'))
    canvas1.create_window(150, 200, window=label)

button1 = tk.Button(
    text='Chạy dữ liệu',
    command=button_run_word, 
    bg='brown',
    fg='white'
    )
canvas1.create_window(button_x, button1_y, window=button1)

# Create "Kiểm tra các file đã được Extract"

def button_check_read_file():
    text = check_read_file()
    string_var.set(text)
    label = tk.Label(inft, textvariable = string_var, fg='black', font=('helvetica', 12, 'bold'))
    canvas1.create_window(150, 200, window=label)

button2 = tk.Button(
    text='Kiểm tra các file đã được đọc',
    command=button_check_read_file, 
    bg='brown',
    fg='white'
    )
canvas1.create_window(button_x, button2_y, window=button2)

# Create "Xóa file xuất để chạy lại dữ liệu"

def button_delete_file():
    text = delete_file()
    string_var.set(text)
    label = tk.Label(inft, textvariable = string_var, fg='black', font=('helvetica', 12, 'bold'))
    canvas1.create_window(150, 200, window=label)
button3 = tk.Button(
    text='Xóa file dữ liệu đã trích xuất',
    command=button_delete_file, 
    bg='brown',
    fg='white'
    )
canvas1.create_window(button_x, button3_y, window=button3)

inft.mainloop()